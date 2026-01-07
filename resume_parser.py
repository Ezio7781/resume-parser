"""
Enhanced Resume Parser with Improved Text Extraction and Parsing Logic
Supports: PDF, DOCX, DOC, TXT with multiple fallback methods
"""

import re
import os
import sys
import json
from typing import List, Dict, Optional, Tuple
import logging
from io import BytesIO
import argparse
import pandas as pd
import concurrent.futures
from functools import partial

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Standard output fields
FIELD_NAMES = [
    'full_name',
    'email',
    'phone_number',
    'alternate_phone_number',
    'highest_qualification',
    'years_of_experience',
    'current_company',
    'current_designation',
    'city',
    'state'
]


# ============================================================================
# TEXT EXTRACTION - ROBUST MULTI-METHOD APPROACH
# ============================================================================

def extract_text(file_path: str) -> str:
    """
    Master text extraction function with intelligent fallbacks.
    Tries multiple methods for each file type until successful.
    """
    if not file_path or not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return ""
    
    ext = os.path.splitext(file_path)[1].lower()
    
    # Route to appropriate handler
    handlers = {
        '.pdf': extract_pdf,
        '.docx': extract_docx,
        '.doc': extract_doc,
        '.txt': extract_txt
    }
    
    handler = handlers.get(ext, extract_txt)
    
    try:
        text = handler(file_path)
        if text and text.strip():
            return clean_text(text)
        
        # Fallback: binary extraction
        logger.warning(f"Primary extraction failed for {file_path}, trying binary fallback")
        return extract_binary_fallback(file_path)
    
    except Exception as e:
        logger.error(f"All extraction methods failed for {file_path}: {e}")
        return ""


def extract_pdf(file_path: str) -> str:
    """Extract text from PDF with 4 fallback methods."""
    text = ""
    
    # Method 1: PyPDF2 (most compatible)
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            logger.debug(f"PyPDF2 extraction successful for {file_path}")
            return text
    except Exception as e:
        logger.debug(f"PyPDF2 failed: {e}")
    
    # Method 2: pdfplumber (better for tables)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            logger.debug(f"pdfplumber extraction successful for {file_path}")
            return text
    except Exception as e:
        logger.debug(f"pdfplumber failed: {e}")
    
    # Method 3: pypdf (alternative)
    try:
        from pypdf import PdfReader
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            logger.debug(f"pypdf extraction successful for {file_path}")
            return text
    except Exception as e:
        logger.debug(f"pypdf failed: {e}")
    
    # Method 4: pdfminer
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(file_path)
        if text and text.strip():
            logger.debug(f"pdfminer extraction successful for {file_path}")
            return text
    except Exception as e:
        logger.debug(f"pdfminer failed: {e}")
    
    return text


def extract_docx(file_path: str) -> str:
    """Extract text from DOCX with 2 methods."""
    # Method 1: python-docx (most reliable)
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        text = "\n".join(text_parts)
        if text.strip():
            logger.debug(f"python-docx extraction successful for {file_path}")
            return text
    except Exception as e:
        logger.debug(f"python-docx failed: {e}")
    
    # Method 2: docx2txt (simpler, faster)
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        if text and text.strip():
            logger.debug(f"docx2txt extraction successful for {file_path}")
            return text
    except Exception as e:
        logger.debug(f"docx2txt failed: {e}")
    
    return ""


def extract_doc(file_path: str) -> str:
    """Extract text from legacy DOC format."""
    # Try converting to DOCX first
    try:
        import subprocess
        import tempfile
        
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "converted.docx")
            
            # Try LibreOffice conversion
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx', 
                 '--outdir', tmpdir, file_path],
                capture_output=True,
                timeout=30
            )
            
            if result.returncode == 0 and os.path.exists(output_path):
                return extract_docx(output_path)
    except Exception as e:
        logger.debug(f"DOC conversion failed: {e}")
    
    # Fallback: try docx library anyway (sometimes works)
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n".join(text_parts)
        if text.strip():
            return text
    except Exception as e:
        logger.debug(f"DOC direct read failed: {e}")
    
    return ""


def extract_txt(file_path: str) -> str:
    """Extract text from TXT with encoding detection."""
    encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'iso-8859-1', 'cp1252', 'utf-16']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                text = f.read()
                if text.strip():
                    logger.debug(f"TXT extraction successful with {encoding}")
                    return text
        except Exception as e:
            logger.debug(f"Failed to read with {encoding}: {e}")
            continue
    
    return ""


def extract_binary_fallback(file_path: str) -> str:
    """Last resort: extract readable text from binary."""
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
        
        # Try each encoding
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = content.decode(encoding, errors='replace')
                # Keep only printable characters
                text = ''.join(c for c in text if c.isprintable() or c in '\n\t')
                if text.strip():
                    return text
            except Exception:
                continue
    except Exception as e:
        logger.error(f"Binary fallback failed: {e}")
    
    return ""


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Remove excessive whitespace
    lines = []
    prev_empty = False
    
    for line in text.split('\n'):
        stripped = line.strip()
        
        if not stripped:
            if not prev_empty:
                lines.append('')
            prev_empty = True
        else:
            # Normalize internal whitespace
            cleaned = ' '.join(stripped.split())
            lines.append(cleaned)
            prev_empty = False
    
    # Remove duplicate consecutive lines
    unique_lines = []
    for line in lines:
        if not unique_lines or unique_lines[-1] != line:
            unique_lines.append(line)
    
    return '\n'.join(unique_lines).strip()


# ============================================================================
# FIELD EXTRACTION - IMPROVED PATTERN MATCHING
# ============================================================================

def extract_name(text: str) -> Optional[str]:
    """Extract candidate name with high confidence."""
    lines = [l.strip() for l in text.split('\n') if l.strip()][:15]
    
    # Blocklist patterns
    blocklist = r'\b(resume|cv|curriculum|vitae|contact|profile|summary|objective|' \
                r'university|college|school|institute|company|inc|corp|llc|ltd|' \
                r'technologies|solutions|systems|email|phone|address)\b'
    
    for line in lines:
        # Skip lines with blocklist keywords
        if re.search(blocklist, line, re.I):
            continue
        
        # Skip lines with emails, phones, URLs
        if re.search(r'@|http|www|\d{7,}', line):
            continue
        
        # Extract name pattern: 2-4 capitalized words
        words = line.split()
        if not (2 <= len(words) <= 4):
            continue
        
# Check each word is a valid name component
        valid = all(re.match(r"^[A-Z][a-z]+(?:['-][A-Z][a-z]+)?$|^[A-Z]\.?$", w) or w in ['Jr', 'Sr', 'II', 'III'] 
                   for w in words)
        
        if valid:
            return line
    
    return None


def extract_email(text: str) -> Optional[str]:
    """Extract email with validation."""
    pattern = r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
    matches = re.findall(pattern, text)
    
    for match in matches:
        email = match.lower().strip()
        # Validate structure
        if '.' in email.split('@')[1] and len(email) >= 6:
            return email
    
    return None


def extract_phone(text: str) -> Optional[str]:
    """Extract primary phone number."""
    # Comprehensive phone patterns
    patterns = [
        r'\+\d{1,3}[\s.-]?\(?\d{1,4}\)?[\s.-]?\d{1,4}[\s.-]?\d{1,9}',  # International
        r'\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',  # US format
        r'\b\d{10,11}\b',  # 10-11 consecutive digits
        r'\d{5}[\s.-]\d{5}',  # Indian format
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            # Normalize: remove non-digits except leading +
            normalized = re.sub(r'[^\d+]', '', match)
            digits = re.sub(r'\D', '', normalized)
            
            # Validate length
            if 10 <= len(digits) <= 15:
                # Add country code if missing
                if len(digits) == 10 and not normalized.startswith('+'):
                    return f"+1{digits}"
                elif not normalized.startswith('+'):
                    return f"+{digits}"
                return normalized
    
    return None


def extract_alternate_phone(text: str) -> Optional[str]:
    """Extract secondary phone number."""
    phones = []
    
    patterns = [
        r'\+\d{1,3}[\s.-]?\(?\d{1,4}\)?[\s.-]?\d{1,4}[\s.-]?\d{1,9}',
        r'\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',
        r'\b\d{10,11}\b',
        r'\d{5}[\s.-]\d{5}',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            normalized = re.sub(r'[^\d+]', '', match)
            digits = re.sub(r'\D', '', normalized)
            
            if 10 <= len(digits) <= 15:
                if len(digits) == 10 and not normalized.startswith('+'):
                    phone = f"+1{digits}"
                elif not normalized.startswith('+'):
                    phone = f"+{digits}"
                else:
                    phone = normalized
                
                if phone not in phones:
                    phones.append(phone)
    
    return phones[1] if len(phones) > 1 else None


def extract_qualification(text: str) -> Optional[str]:
    """Extract highest qualification from Education section."""
    # Find education section
    edu_section = extract_section(text, ['education', 'academic', 'qualifications'])
    if not edu_section:
        edu_section = text
    
    # Degree mappings (highest to lowest)
    degrees = [
        ('PhD', ['phd', 'ph.d', 'ph d', 'doctorate', 'doctor of philosophy', 'doctoral']),
        ('Masters', ['master', 'masters', 'mba', 'ms', 'm.s', 'msc', 'm.sc', 'mtech', 'm.tech', 'ma', 'm.a']),
        ('Bachelors', ['bachelor', 'bachelors', 'btech', 'b.tech', 'be', 'b.e', 'bs', 'b.s', 'bsc', 'b.sc', 'ba', 'b.a', 'bcom', 'b.com']),
        ('Diploma', ['diploma', 'dip', 'polytechnic']),
        ('Associate', ['associate', 'assoc'])
    ]
    
    edu_lower = edu_section.lower()
    
    for degree_name, keywords in degrees:
        for kw in keywords:
            # Use word boundaries to avoid false matches
            if re.search(rf'\b{re.escape(kw)}\b', edu_lower):
                return degree_name
    
    return None


def extract_experience(text: str) -> Optional[float]:
    """Extract years of experience."""
    # Priority patterns (most specific first)
    patterns = [
        r'(\d+(?:\.\d+)?)\s*\+?\s*years?\s+of\s+(?:work|professional|total)?\s*experience',
        r'experience\s*[:\-]\s*(\d+(?:\.\d+)?)\s*\+?\s*years?',
        r'(\d+(?:\.\d+)?)\s*\+?\s*(?:yrs?|years?)\s+experience',
        r'total\s+experience\s*[:\-]?\s*(\d+(?:\.\d+)?)\s*years?',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.I)
        for match in matches:
            try:
                years = float(match.group(1))
                if 0 < years <= 50:  # Reasonable range
                    return years
            except (ValueError, IndexError):
                continue
    
    return None


def extract_current_company(text: str) -> Optional[str]:
    """Extract current/most recent company."""
    exp_section = extract_section(text, ['experience', 'work experience', 'employment', 'professional'])
    
    if not exp_section:
        exp_section = text  # Fallback to full text
    
    lines = [l.strip() for l in exp_section.split('\n') if l.strip()]
    
    # Look for "Present" indicator and find associated company
    for i, line in enumerate(lines[:30]):
        if re.search(r'\b(present|current|till\s+now|ongoing)\b', line, re.I):
            # Check surrounding lines for company name
            for offset in range(-3, 4):
                idx = i + offset
                if 0 <= idx < len(lines):
                    candidate = lines[idx]
                    
                    # Clean date patterns
                    candidate = re.sub(r'\b\w{3,4}\s+\d{4}\s*[-–]\s*(present|current|\d{4})\b', '', candidate, flags=re.I)
                    candidate = candidate.strip()
                    
                    # Skip if it's a job title or contains title keywords
                    title_keywords = r'\b(engineer|developer|analyst|manager|director|lead|senior|junior|' \
                                   r'associate|specialist|consultant|architect|designer|coordinator|' \
                                   r'executive|officer|administrator|supervisor|intern|trainee)\b'
                    
                    if not re.search(title_keywords, candidate, re.I) and _looks_like_company(candidate):
                        return candidate
    
    # Fallback: Look for "at Company" patterns
    for line in lines[:20]:
        at_match = re.search(r'(.+?)\s+at\s+(.+)', line, re.I)
        if at_match:
            title_part = at_match.group(1).strip()
            company_part = at_match.group(2).strip()
            
            # Check if title part has title keywords
            title_keywords = r'\b(engineer|developer|analyst|manager|director|lead|senior|junior|' \
                           r'associate|specialist|consultant|architect|designer|coordinator|' \
                           r'executive|officer|administrator|supervisor|intern|trainee)\b'
            
            if re.search(title_keywords, title_part, re.I) and _looks_like_company(company_part):
                return company_part
    
    # Another fallback: Look for well-known company names
    known_companies = ['google', 'microsoft', 'amazon', 'apple', 'facebook', 'meta', 'netflix', 
                     'tesla', 'ibm', 'oracle', 'adobe', 'salesforce', 'linkedin', 'twitter']
    
    text_lower = text.lower()
    for company in known_companies:
        if re.search(rf'\b{re.escape(company)}\b', text_lower):
            return company.title()
    
    return None


def extract_designation(text: str) -> Optional[str]:
    """Extract current/most recent job title."""
    exp_section = extract_section(text, ['experience', 'work experience', 'employment'])
    
    if not exp_section:
        exp_section = text  # Fallback to full text
    
    lines = [l.strip() for l in exp_section.split('\n') if l.strip()]
    
    # Job title keywords
    title_keywords = r'\b(engineer|developer|analyst|manager|director|lead|senior|junior|' \
                     r'associate|specialist|consultant|architect|designer|coordinator|' \
                     r'executive|officer|administrator|supervisor|intern|trainee)\b'
    
    # Look for "Present" context
    for i, line in enumerate(lines[:30]):
        if re.search(r'\b(present|current)\b', line, re.I):
            # Check surrounding lines for job title
            for offset in range(-3, 4):
                idx = i + offset
                if 0 <= idx < len(lines):
                    candidate = lines[idx]
                    
                    # Clean date patterns
                    candidate = re.sub(r'\b\w{3,4}\s+\d{4}\s*[-–]\s*(present|current|\d{4})\b', '', candidate, flags=re.I)
                    candidate = candidate.strip()
                    
                    # Skip if looks like company name
                    if not _looks_like_company(candidate) and re.search(title_keywords, candidate, re.I):
                        # Extract just the title part
                        title = _extract_title_from_line(candidate)
                        if title and len(title) < 80:
                            return title
    
    # Fallback: Look for "Title at Company" patterns
    for line in lines[:20]:
        at_match = re.search(r'(.+?)\s+at\s+(.+)', line, re.I)
        if at_match:
            title_part = at_match.group(1).strip()
            company_part = at_match.group(2).strip()
            
            # Check if title part has title keywords and company part looks like company
            if re.search(title_keywords, title_part, re.I) and _looks_like_company(company_part):
                title = _extract_title_from_line(title_part)
                if title and len(title) < 80:
                    return title
    
    # Another fallback: Find first line with job title keyword
    for line in lines[:15]:
        if not _looks_like_company(line) and re.search(title_keywords, line, re.I):
            title = _extract_title_from_line(line)
            if title and len(title) < 80:
                return title
    
    return None


def extract_city(text: str) -> Optional[str]:
    """Extract city name."""
    # Exclude skills section
    text_lower = text.lower()
    skills_pos = text_lower.find('skills')
    if skills_pos != -1:
        text_to_search = text[:skills_pos]
    else:
        text_to_search = text
    
    # Common cities (expandable)
    cities = [
        'bangalore', 'bengaluru', 'mumbai', 'delhi', 'hyderabad', 'chennai', 
        'kolkata', 'pune', 'ahmedabad', 'jaipur', 'surat', 'lucknow', 'kanpur',
        'new york', 'san francisco', 'los angeles', 'chicago', 'boston', 'seattle',
        'london', 'paris', 'berlin', 'toronto', 'sydney', 'singapore'
    ]
    
    for city in cities:
        if re.search(rf'\b{re.escape(city)}\b', text_to_search.lower()):
            return city.title()
    
    # Pattern: City, State
    location_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z][a-z]+)'
    matches = re.findall(location_pattern, text_to_search)
    if matches:
        return matches[0][0]
    
    return None


def extract_state(text: str) -> Optional[str]:
    """Extract state/province name."""
    # Exclude skills section
    text_lower = text.lower()
    skills_pos = text_lower.find('skills')
    if skills_pos != -1:
        text_to_search = text[:skills_pos]
    else:
        text_to_search = text
    
    # Common states (expandable)
    states = [
        'california', 'new york', 'texas', 'florida', 'illinois', 'washington',
        'maharashtra', 'karnataka', 'tamil nadu', 'delhi', 'uttar pradesh',
        'west bengal', 'gujarat', 'rajasthan', 'telangana', 'andhra pradesh'
    ]
    
    for state in states:
        if re.search(rf'\b{re.escape(state)}\b', text_to_search.lower()):
            return state.title()
    
    # Pattern: City, State
    location_pattern = r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)'
    matches = re.findall(location_pattern, text_to_search)
    if matches:
        return matches[0]
    
    return None


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def extract_section(text: str, headers: List[str], window: int = 1500) -> Optional[str]:
    """Extract text section by headers."""
    lines = text.split('\n')
    text_lower = text.lower()
    
    for header in headers:
        # Look for header as a standalone line (or with minimal following text)
        for i, line in enumerate(lines):
            line_lower = line.strip().lower()
            
            # Check if line starts with header (exact match or followed by colon/space)
            if (line_lower == header or 
                line_lower.startswith(header + ':') or 
                line_lower.startswith(header + ' ')):
                
                # Get the position in original text
                line_start_pos = text.find(line)
                
                # Find end of this line (start after header line)
                line_end = text.find('\n', line_start_pos)
                if line_end == -1:
                    line_end = len(text)
                
                # Start after this line
                section_start = line_end + 1
                
                # Get remaining text starting from after header line
                remaining_text = text[section_start:section_start + window]
                
                # Try to find next section header
                next_headers = ['education', 'experience', 'skills', 'projects', 'certification', 'summary', 'objective', 'contact']
                min_next = len(remaining_text)
                
                for nh in next_headers:
                    if nh != header:
                        # Look for next header as standalone line
                        next_pos = -1
                        next_lines = remaining_text.split('\n')
                        for j, next_line in enumerate(next_lines):
                            next_line_lower = next_line.strip().lower()
                            if (next_line_lower == nh or 
                                next_line_lower.startswith(nh + ':') or 
                                next_line_lower.startswith(nh + ' ')):
                                next_pos = remaining_text.find(next_line)
                                break
                        
                        if next_pos != -1 and next_pos < min_next:
                            min_next = next_pos
                
                return remaining_text[:min_next].strip()
    
    return None


def _looks_like_company(text: str) -> bool:
    """Check if text looks like a company name."""
    if not text or len(text) < 2 or len(text) > 100:
        return False
    
    # Company indicators
    indicators = r'\b(inc|corp|corporation|company|ltd|llc|co|limited|technologies|' \
                 r'solutions|systems|group|services|consulting|pvt)\b'
    
    if re.search(indicators, text, re.I):
        return True
    
    # Check if starts with capital and has reasonable length
    words = text.split()
    if 1 <= len(words) <= 6 and text[0].isupper():
        return True
    
    return False


def _extract_title_from_line(line: str) -> Optional[str]:
    """Extract job title from line containing title and company."""
    # Try common separators
    separators = [' at ', ' | ', ' - ', ' — ', ' – ']
    
    for sep in separators:
        if sep in line:
            parts = line.split(sep)
            # First part is usually the title
            title = parts[0].strip()
            if len(title) < 80 and title:
                return title
    
    # If no separator, return cleaned line if not too long
    if len(line) < 80:
        # Remove date patterns
        cleaned = re.sub(r'\b\w{3}\s+\d{4}\s*[-–]\s*\w+\b', '', line)
        cleaned = cleaned.strip()
        if cleaned:
            return cleaned
    
    return None


# ============================================================================
# MAIN PARSING FUNCTION
# ============================================================================

def parse_resume(file_path: str) -> Dict[str, Optional[object]]:
    """
    Main parsing function - combines rule-based extraction.
    """
    # Extract text
    text = extract_text(file_path)
    
    if not text:
        logger.error(f"Failed to extract text from {file_path}")
        return {field: None for field in [
            'full_name', 'email', 'phone_number', 'alternate_phone_number',
            'highest_qualification', 'years_of_experience', 'current_company',
            'current_designation', 'city', 'state'
        ]}
    
    # Rule-based extraction
    result = {
        'full_name': extract_name(text),
        'email': extract_email(text),
        'phone_number': extract_phone(text),
        'alternate_phone_number': extract_alternate_phone(text),
        'highest_qualification': extract_qualification(text),
        'years_of_experience': extract_experience(text),
        'current_company': extract_current_company(text),
        'current_designation': extract_designation(text),
        'city': extract_city(text),
        'state': extract_state(text)
    }
    
    return result


# CLI interface
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Parse resume files')
    parser.add_argument('file_path', help='Path to resume file')
    parser.add_argument('--output', '-o', help='Output JSON file')
    
    args = parser.parse_args()
    
    result = parse_resume(args.file_path)
    
    output = json.dumps(result, indent=2)
    
    if args.output:
        with open(args.output, 'w') as f:
            f.write(output)
        print(f"Results written to {args.output}")
    else:
        print(output)
