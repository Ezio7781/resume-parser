from flask import Flask, request, jsonify, send_file, render_template_string, url_for, abort
import os
import re
from datetime import datetime
import pandas as pd
from io import BytesIO
import json
import time
import threading
import shutil
import concurrent.futures
import docx2txt
from functools import wraps

try:
    from secrets_store import store_api_key, get_api_key as get_stored_api_key, delete_api_key
except Exception:
    store_api_key = None
    get_stored_api_key = None
    delete_api_key = None

try:
    from llm_helper import call_llm_extract
except Exception:
    call_llm_extract = None

try:
    from resume_parser import parse_resume, extract_text
except Exception:
    parse_resume = None
    extract_text = None


app = Flask(__name__)

# Security Configuration
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('PARSE_MAX_FILE_MB', '5')) * 1024 * 1024
app.config['SESSION_COOKIE_SECURE'] = os.getenv('FLASK_ENV', 'production') != 'development'
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# Security Headers Middleware
@app.after_request
def add_security_headers(response):
    """Add security headers to every response"""
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    response.headers['Content-Security-Policy'] = "default-src 'self'; script-src 'self' 'unsafe-inline'; style-src 'self' 'unsafe-inline'; img-src 'self' data:; font-src 'self'"
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Permissions-Policy'] = 'geolocation=(), microphone=(), camera=()'
    return response

HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume Parser AI Agent</title>
    <style>
        {% if not default_theme %}
        :root {
            --primary: #3B82F6; /* blue */
            --primary-dark: #1E40AF;
            --secondary: #8B5CF6; /* purple */
            --success: #10B981;
            --success-dark: #059669;
            --danger: #EF4444;
            --warning: #F59E0B;
            --bg-primary: #F9FAFB;
            --bg-secondary: #F3F4F6;
            --bg-tertiary: #E5E7EB;
            --text-primary: #111827;
            --text-secondary: #374151;
            --text-muted: #6B7280;
            --border-color: #D1D5DB;
            --shadow: rgba(0, 0, 0, 0.05);
            --shadow-lg: rgba(0, 0, 0, 0.1);
        }

        [data-theme="dark"] {
            --primary: #818cf8;
            --primary-dark: #6366f1;
            --secondary: #a78bfa;
            --success: #34d399;
            --success-dark: #10b981;
            --danger: #f87171;
            --warning: #fbbf24;
            --bg-primary: #1e293b;
            --bg-secondary: #0f172a;
            --bg-tertiary: #334155;
            --text-primary: #f1f5f9;
            --text-secondary: #cbd5e1;
            --text-muted: #94a3b8;
            --border-color: #334155;
            --shadow: rgba(0, 0, 0, 0.3);
            --shadow-lg: rgba(0, 0, 0, 0.5);
        }
        {% else %}
        /* Server-enforced single theme to reduce payload */
        :root {
            {% if default_theme == 'dark' %}
            --primary: #818cf8;
            --primary-dark: #6366f1;
            --secondary: #a78bfa;
            --success: #34d399;
            --success-dark: #10b981;
            --danger: #f87171;
            --warning: #fbbf24;
            --bg-primary: #1e293b;
            --bg-secondary: #0f172a;
            --bg-tertiary: #334155;
            --text-primary: #f1f5f9;
            --text-secondary: #cbd5e1;
            --text-muted: #94a3b8;
            --border-color: #334155;
            --shadow: rgba(0, 0, 0, 0.3);
            --shadow-lg: rgba(0, 0, 0, 0.5);
            {% else %}
            --primary: #3B82F6;
            --primary-dark: #1E40AF;
            --secondary: #8B5CF6;
            --success: #10B981;
            --success-dark: #059669;
            --danger: #EF4444;
            --warning: #F59E0B;
            --bg-primary: #F9FAFB;
            --bg-secondary: #F3F4F6;
            --bg-tertiary: #E5E7EB;
            --text-primary: #111827;
            --text-secondary: #374151;
            --text-muted: #6B7280;
            --border-color: #D1D5DB;
            --shadow: rgba(0, 0, 0, 0.05);
            --shadow-lg: rgba(0, 0, 0, 0.1);
            {% endif %}
        }
        {% endif %}

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
            transition: background-color 0.3s ease, color 0.3s ease, border-color 0.3s ease;
        }
        
        *::selection {
            background: transparent;
            color: inherit;
        }
        
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
            background: var(--bg-secondary);
            min-height: 100vh;
            padding: 20px;
            color: var(--text-primary);
            scroll-behavior: smooth;
            overflow-x: hidden;
        }
        
        .container {
            max-width: 1400px;
            margin: 0 auto;
        }

        .header {
            background: linear-gradient(135deg, var(--primary) 0%, var(--secondary) 100%);
            border-radius: 20px;
            box-shadow: 0 20px 60px var(--shadow-lg);
            padding: 40px;
            margin-bottom: 30px;
            position: relative;
            overflow: hidden;
        }

        .header::before {
            content: '';
            position: absolute;
            top: -50%;
            right: -50%;
            width: 200%;
            height: 200%;
            background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
            animation: float 15s infinite ease-in-out;
        }

        @keyframes float {
            0%, 100% { transform: translate(0, 0) rotate(0deg); }
            50% { transform: translate(-20px, -20px) rotate(180deg); }
        }

        @keyframes rotate {
            from { transform: rotate(0deg); }
            to { transform: rotate(360deg); }
        }

        @keyframes slideInDown {
            from {
                opacity: 0;
                transform: translateY(-20px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }

        @keyframes slideInUp {
            from {
                opacity: 0;
                transform: translateY(20px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }

        @keyframes pulse {
            0%, 100% { opacity: 1; }
            50% { opacity: 0.5; }
        }

        @keyframes shimmer {
            0% { background-position: -1000px 0; }
            100% { background-position: 1000px 0; }
        }

        .header-content {
            position: relative;
            z-index: 1;
        }

        .header-top {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
        }

        h1 {
            font-size: 2.5em;
            color: white;
            display: flex;
            align-items: center;
            gap: 15px;
            margin: 0;
        }

        {% if not default_theme %}
        .theme-toggle {
            background: rgba(255, 255, 255, 0.12);
            border: 1px solid rgba(255, 255, 255, 0.18);
            color: white;
            padding: 10px 18px;
            border-radius: 12px;
            cursor: pointer;
            font-size: 1.05em;
            backdrop-filter: blur(6px);
            transition: all 0.2s ease;
            box-shadow: 0 6px 18px rgba(79,70,229,0.18);
        }
        {% endif %}

        .snow-toggle {
            background: rgba(255, 255, 255, 0.12);
            border: 1px solid rgba(255, 255, 255, 0.14);
            color: white;
            padding: 8px 12px;
            border-radius: 12px;
            cursor: pointer;
            font-size: 1.05em;
            margin-left: 10px;
            backdrop-filter: blur(6px);
            transition: all 0.2s ease;
            box-shadow: 0 6px 18px rgba(6,182,212,0.12);
        }

        /* full-screen snow canvas */
        #snowCanvas {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            pointer-events: none;
            z-index: 9999;
            opacity: 0.95;
        }

        {% if not default_theme %}
        .theme-toggle:hover {
            background: rgba(255, 255, 255, 0.3);
            transform: scale(1.05);
        }
        {% endif %}
        
        .subtitle {
            color: rgba(255, 255, 255, 0.9);
            font-size: 1.1em;
        }

        .stats-container {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }

        .stat-card {
            background: var(--bg-primary);
            border-radius: 15px;
            padding: 25px;
            box-shadow: 0 4px 15px var(--shadow);
            border: 1px solid var(--border-color);
        }

        .stat-label {
            color: var(--text-secondary);
            font-size: 0.9em;
            margin-bottom: 8px;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }

        .stat-value {
            font-size: 2em;
            font-weight: bold;
            color: var(--text-primary);
        }

        .stat-value.success { color: var(--success); }
        .stat-value.danger { color: var(--danger); }
        .stat-value.warning { color: var(--warning); }

        .main-card {
            background: var(--bg-primary);
            border-radius: 20px;
            box-shadow: 0 10px 40px var(--shadow);
            padding: 40px;
            margin-bottom: 30px;
            border: 1px solid var(--border-color);
        }
        
        .upload-area {
            border: 3px dashed var(--border-color);
            border-radius: 15px;
            padding: 60px 40px;
            text-align: center;
            margin: 30px 0;
            background: var(--bg-secondary);
            transition: all 0.4s cubic-bezier(0.34, 1.56, 0.64, 1);
            cursor: pointer;
            position: relative;
            overflow: hidden;
        }

        .upload-area::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(90deg, transparent, rgba(255,255,255,0.1), transparent);
            transition: left 0.5s ease;
        }
        
        .upload-area:hover {
            border-color: var(--primary);
            background: var(--bg-tertiary);
            transform: translateY(-4px) scale(1.01);
            box-shadow: 0 10px 30px rgba(217, 119, 6, 0.2);
        }

        .upload-area:hover::before {
            left: 100%;
        }
        
        .upload-area.dragover {
            border-color: var(--success);
            background: var(--bg-tertiary);
            transform: scale(1.04);
            box-shadow: 0 15px 40px rgba(22, 163, 74, 0.25);
            border-width: 4px;
        }
        
        .upload-icon {
            font-size: 4em;
            margin-bottom: 15px;
            animation: bounce 2s infinite;
            display: inline-block;
        }

        @keyframes bounce {
            0%, 100% { transform: translateY(0) scale(1); }
            50% { transform: translateY(-15px) scale(1.05); }
        }
        
        .upload-text {
            font-size: 1.2em;
            color: var(--text-primary);
            margin-bottom: 10px;
            font-weight: 600;
        }
        
        .upload-hint {
            color: var(--text-muted);
            font-size: 0.95em;
        }

        /* Modal / settings */
        .modal-backdrop { position: fixed; inset: 0; background: rgba(2,6,23,0.6); display: none; align-items: center; justify-content: center; z-index: 10001; }
        .modal { background: var(--bg-primary); width: 720px; max-width: 94%; border-radius: 12px; padding: 20px; box-shadow: 0 20px 60px var(--shadow-lg); }
        .modal-header { display:flex; justify-content:space-between; align-items:center; margin-bottom:12px }
        .modal-title { font-weight:700; font-size:1.1em }
        .modal-body { max-height: 60vh; overflow:auto; }
        .row { display:flex; gap:12px; margin-bottom:12px; align-items:center }
        .field { flex:1 }
        .small { font-size:0.9em; color:var(--text-muted) }
        .pill { padding:8px 12px; border-radius:999px; background:var(--bg-secondary); border:1px solid var(--border-color); }
        
        #fileInput {
            display: none;
        }
        
        .file-list {
            margin: 20px 0;
            padding: 20px;
            background: var(--bg-secondary);
            border-radius: 10px;
            display: none;
        }

        .file-list-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 15px;
        }

        .file-list-title {
            font-size: 1.1em;
            font-weight: 600;
            color: var(--text-primary);
        }

        .clear-files {
            background: var(--danger);
            color: white;
            border: none;
            padding: 8px 16px;
            border-radius: 8px;
            cursor: pointer;
            font-size: 0.9em;
            transition: all 0.3s ease;
        }

        .clear-files:hover {
            background: #dc2626;
            transform: translateY(-2px);
        }
        
        .file-item {
            display: flex;
            align-items: center;
            gap: 12px;
            padding: 12px;
            background: var(--bg-primary);
            margin-bottom: 8px;
            border-radius: 8px;
            border-left: 4px solid var(--primary);
            position: relative;
            overflow: hidden;
        }

        .file-item.parsing {
            border-left-color: var(--warning);
        }

        .file-item.success {
            border-left-color: var(--success);
        }

        .file-item.failed {
            border-left-color: var(--danger);
        }

        .file-icon {
            font-size: 1.5em;
        }

        .file-info {
            flex: 1;
        }

        .file-name {
            color: var(--text-primary);
            font-weight: 500;
        }

        .file-size {
            color: var(--text-muted);
            font-size: 0.85em;
        }

        .file-status {
            display: flex;
            align-items: center;
            gap: 8px;
            font-size: 0.9em;
            font-weight: 600;
        }

        .file-status.pending { color: var(--text-muted); }
        .file-status.parsing { color: var(--warning); }
        .file-status.success { color: var(--success); }
        .file-status.failed { color: var(--danger); }

        .file-progress {
            position: absolute;
            bottom: 0;
            left: 0;
            height: 3px;
            background: var(--primary);
            width: 0%;
            transition: width 0.3s ease;
        }
        
        .btn-container {
            display: flex;
            gap: 15px;
            margin: 20px 0;
        }

        .btn {
            flex: 1;
            background: linear-gradient(135deg, var(--primary) 0%, var(--secondary) 100%);
            color: white;
            padding: 15px 40px;
            border: none;
            border-radius: 10px;
            font-size: 1.1em;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 10px;
        }
        
        .btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
        }
        
        .btn:disabled {
            background: var(--text-muted);
            cursor: not-allowed;
            transform: none;
            box-shadow: none;
        }
        
        .btn-success {
            background: linear-gradient(135deg, var(--success) 0%, var(--success-dark) 100%);
            box-shadow: 0 4px 15px rgba(16, 185, 129, 0.4);
        }
        
        .btn-success:hover {
            box-shadow: 0 6px 20px rgba(16, 185, 129, 0.6);
        }

        .btn-secondary {
            background: var(--bg-tertiary);
            color: var(--text-primary);
            box-shadow: 0 2px 10px var(--shadow);
        }

        .btn-secondary:hover {
            background: var(--border-color);
        }
        
        .progress-container {
            display: none;
            margin: 30px 0;
        }

        .progress-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 15px;
        }

        .progress-title {
            font-size: 1.2em;
            font-weight: 600;
            color: var(--text-primary);
        }

        .progress-stats {
            display: flex;
            gap: 20px;
            font-size: 0.9em;
        }

        .progress-stat {
            display: flex;
            align-items: center;
            gap: 5px;
        }

        .progress-bar-container {
            background: var(--bg-secondary);
            border-radius: 10px;
            height: 30px;
            overflow: hidden;
            position: relative;
            box-shadow: inset 0 2px 4px var(--shadow);
        }

        .progress-bar {
            height: 100%;
            background: linear-gradient(90deg, var(--primary), var(--secondary));
            transition: width 0.5s ease;
            display: flex;
            align-items: center;
            justify-content: center;
            color: white;
            font-weight: 600;
            font-size: 0.9em;
        }

        .progress-bar.complete {
            background: linear-gradient(90deg, var(--success), var(--success-dark));
        }
        
        .alert {
            padding: 15px 20px;
            border-radius: 10px;
            margin: 20px 0;
            display: none;
            animation: slideIn 0.3s ease;
        }

        @keyframes slideIn {
            from { opacity: 0; transform: translateY(-10px); }
            to { opacity: 1; transform: translateY(0); }
        }
        
        .alert-success {
            background: rgba(16, 185, 129, 0.1);
            color: var(--success);
            border-left: 4px solid var(--success);
        }
        
        .alert-error {
            background: rgba(239, 68, 68, 0.1);
            color: var(--danger);
            border-left: 4px solid var(--danger);
        }

        .alert-warning {
            background: rgba(245, 158, 11, 0.1);
            color: var(--warning);
            border-left: 4px solid var(--warning);
        }
        
        .results-container {
            margin-top: 30px;
            display: none;
        }
        
        .results-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
            padding-bottom: 15px;
            border-bottom: 2px solid var(--border-color);
            flex-wrap: wrap;
            gap: 15px;
        }
        
        .results-title {
            font-size: 1.5em;
            color: var(--text-primary);
            font-weight: 600;
        }

        .results-actions {
            display: flex;
            gap: 10px;
        }

        .action-btn {
            background: var(--bg-secondary);
            color: var(--text-primary);
            border: 1px solid var(--border-color);
            padding: 10px 20px;
            border-radius: 8px;
            cursor: pointer;
            font-size: 0.9em;
            transition: all 0.3s ease;
            display: flex;
            align-items: center;
            gap: 8px;
        }

        .action-btn:hover {
            background: var(--bg-tertiary);
            transform: translateY(-2px);
        }

        .search-box {
            padding: 10px 15px;
            border: 1px solid var(--border-color);
            border-radius: 8px;
            background: var(--bg-secondary);
            color: var(--text-primary);
            font-size: 0.95em;
            width: 250px;
        }

        .search-box:focus {
            outline: none;
            border-color: var(--primary);
        }
        
        .table-responsive {
            overflow-x: auto;
            border-radius: 10px;
            box-shadow: 0 2px 10px var(--shadow);
        }

        table {
            width: 100%;
            border-collapse: collapse;
            background: var(--bg-primary);
        }
        
        thead {
            background: linear-gradient(135deg, var(--primary) 0%, var(--secondary) 100%);
            color: white;
            position: sticky;
            top: 0;
            z-index: 10;
        }
        
        th {
            padding: 15px;
            text-align: left;
            font-weight: 600;
            font-size: 0.9em;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            white-space: nowrap;
        }
        
        td {
            padding: 12px 15px;
            border-bottom: 1px solid var(--border-color);
            color: var(--text-primary);
        }
        
        tbody tr {
            transition: all 0.3s ease;
            background: var(--bg-primary);
        }
        
        tbody tr:hover {
            background: var(--bg-secondary);
            transform: translateX(4px);
            box-shadow: -4px 0 12px rgba(217, 119, 6, 0.15);
        }
        
        tbody tr:nth-child(even) {
            background: rgba(0, 0, 0, 0.02);
        }
        
        [data-theme="dark"] tbody tr:nth-child(even) {
            background: rgba(255, 255, 255, 0.02);
        }
            transition: background 0.2s ease;
        }
        
        tbody tr:hover {
            background: var(--bg-secondary);
        }
        
        .null-value {
            color: var(--text-muted);
            font-style: italic;
        }

        .filter-tabs {
            display: flex;
            gap: 10px;
            margin-bottom: 20px;
            flex-wrap: wrap;
        }

        .filter-tab {
            padding: 8px 16px;
            border: 2px solid var(--border-color);
            background: var(--bg-secondary);
            color: var(--text-primary);
            border-radius: 20px;
            cursor: pointer;
            transition: all 0.3s ease;
            font-size: 0.9em;
            font-weight: 500;
        }

        .filter-tab:hover {
            border-color: var(--primary);
        }

        .filter-tab.active {
            background: var(--primary);
            color: white;
            border-color: var(--primary);
        }

        .empty-state {
            text-align: center;
            padding: 60px 20px;
            color: var(--text-muted);
        }

        .empty-icon {
            font-size: 4em;
            margin-bottom: 20px;
            opacity: 0.5;
        }

        @media (max-width: 768px) {
            h1 { font-size: 1.8em; }
            .stats-container { grid-template-columns: 1fr 1fr; }
            .btn-container { flex-direction: column; }
            .results-header { flex-direction: column; align-items: flex-start; }
            .search-box { width: 100%; }
        }
    </style>
</head>
<body data-theme="{{ default_theme or 'light' }}" data-server-default-theme="{{ default_theme }}">
    <canvas id="snowCanvas"></canvas>
    <div class="container">
        <div class="header">
            <div class="header-content">
                <div class="header-top">
                    <h1>📄 Resume Parser AI Agent</h1>
                    <div style="display:flex; align-items:center; gap:8px;">
                        {% if not default_theme %}
                        <button class="theme-toggle" onclick="toggleTheme()" title="Toggle dark/light">🌙</button>
                        {% endif %}
                        <button class="snow-toggle" id="snowToggle" onclick="toggleSnow()" title="Toggle snow">❄️ Snow</button>
                        <button class="snow-toggle" id="settingsBtn" onclick="openSettings()" title="Settings">⚙️</button>
                        <button class="snow-toggle" id="clearParsedBtn" onclick="clearParsed()" title="Clear parsed previews">🧹 Clear Parsed</button>
                    </div>
                </div>
                <p class="subtitle">Upload multiple resumes and extract structured candidate information with real-time progress tracking</p>
            </div>
        </div>

        <div class="stats-container">
            <div class="stat-card">
                <div class="stat-label">Total Uploaded</div>
                <div class="stat-value" id="totalFiles">0</div>
            </div>
            <div class="stat-card">
                <div class="stat-label">Successfully Parsed</div>
                <div class="stat-value success" id="successCount">0</div>
            </div>
            <div class="stat-card">
                <div class="stat-label">Failed</div>
                <div class="stat-value danger" id="failedCount">0</div>
            </div>
            <div class="stat-card">
                <div class="stat-label">Processing Time</div>
                <div class="stat-value warning" id="processingTime">0s</div>
            </div>
        </div>

        <div class="main-card">
            <div class="upload-area" id="uploadArea" onclick="document.getElementById('fileInput').click()">
                <div class="upload-icon">📤</div>
                <div class="upload-text">Click to upload or drag and drop resumes</div>
                <div class="upload-hint">Supports: PDF, DOC, DOCX, TXT | Max: {{max_uploads}} files | Max file: {{max_file_size_mb}} MB</div>
                <input type="file" id="fileInput" multiple accept=".txt,.pdf,.doc,.docx">
            </div>
            
            <div class="file-list" id="fileList"></div>
            
            <div class="btn-container">
                <button class="btn" id="parseBtn" onclick="parseResumes()" disabled title="Parse selected resumes">
                    <span id="parseIcon">🚀</span>&nbsp;<span id="parseLabel">Parse Resumes</span>
                </button>
                <button class="btn btn-secondary" id="clearBtn" onclick="clearAll()" style="display: none;">
                    🗑️ Clear All
                </button>
            </div>

            <div class="progress-container" id="progressContainer">
                <div class="progress-header">
                    <div class="progress-title">⚡ Parsing Progress</div>
                    <div class="progress-stats">
                        <div class="progress-stat">
                            <span>📊</span>
                            <span id="progressText">0/0</span>
                        </div>
                    </div>
                </div>
                <div class="progress-bar-container">
                    <div class="progress-bar" id="progressBar">0%</div>
                </div>
            </div>
            
            <div class="alert alert-success" id="successAlert"></div>
            <div class="alert alert-error" id="errorAlert"></div>
            <div class="alert alert-warning" id="warningAlert"></div>
        </div>

        <!-- Settings / Controls Modal -->
        <div class="modal-backdrop" id="modalBackdrop">
            <div class="modal" role="dialog" aria-modal="true" aria-labelledby="modalTitle">
                <div class="modal-header">
                    <div class="modal-title" id="modalTitle">Settings</div>
                    <div><button class="pill" onclick="closeSettings()">Close</button></div>
                </div>
                <div class="modal-body">
                    <div class="row">
                        <div class="field">
                            <label><strong>Enable LLM Extraction</strong></label>
                            <div class="small">When enabled, extracted fields use the configured LLM. API key stored in memory only.</div>
                        </div>
                        <div style="width:160px; text-align:right">
                            <label class="pill"><input type="checkbox" id="llmToggle"> Use LLM</label>
                        </div>
                    </div>
                    <div class="row">
                        <div class="field">
                            <label><strong>Model</strong></label>
                            <select id="modelSelect" class="field">
                                <option value="gpt-4o-mini">gpt-4o-mini</option>
                                <option value="gpt-4">gpt-4</option>
                                <option value="grok-1">grok-1</option>
                            </select>
                        </div>
                        <div style="width:220px">
                            <label><strong>Snow Intensity</strong></label>
                            <input type="range" id="snowIntensity" min="0" max="200" value="120" style="width:100%">
                        </div>
                    </div>
                    <div class="row">
                        <div class="field">
                            <label><strong>API Key (optional)</strong></label>
                            <input id="apiKeyInput" type="password" placeholder="Enter API key for LLM calls" style="width:100%; padding:8px; border-radius:8px; border:1px solid var(--border-color)">
                            <div class="small">Key is used only for extraction requests from the browser to the server.</div>
                        </div>
                    </div>
                    <div class="row">
                        <button class="btn" onclick="saveSettings()">Save Settings</button>
                        <button class="btn btn-secondary" onclick="downloadSample()">Download Sample Resume</button>
                        <button class="btn btn-secondary" onclick="clearLocalSettings()">Reset</button>
                    </div>
                </div>
            </div>
        </div>

        <!-- Preview Modal -->
        <div class="modal-backdrop" id="previewBackdrop" style="display:none">
            <div class="modal">
                <div class="modal-header">
                    <div class="modal-title">Parsed JSON Preview</div>
                    <div><button class="pill" onclick="closePreview()">Close</button></div>
                </div>
                <div class="modal-body"><pre id="previewContent" style="white-space:pre-wrap"></pre></div>
            </div>
        </div>

        <div class="results-container" id="resultsContainer">
            <div class="main-card">
                <div class="results-header">
                    <div class="results-title">📊 Parsed Results</div>
                    <div class="results-actions">
                        <input type="text" class="search-box" id="searchBox" placeholder="🔍 Search candidates..." onkeyup="searchTable()">
                        <button class="action-btn" onclick="downloadExcel()">
                            💾 Download Excel
                        </button>
                        <button class="action-btn" onclick="downloadJSON()">
                            📋 Download JSON
                        </button>
                    </div>
                </div>

                <div class="filter-tabs" id="filterTabs">
                    <div class="filter-tab active" onclick="filterResults('all')">All Candidates</div>
                    <div class="filter-tab" onclick="filterResults('success')">✅ Successful</div>
                    <div class="filter-tab" onclick="filterResults('failed')">❌ Failed</div>
                    <div class="filter-tab" onclick="filterResults('incomplete')">⚠️ Incomplete Data</div>
                </div>

                <div class="table-responsive">
                    <table id="resultsTable"></table>
                </div>
            </div>
        </div>
    </div>
    
    <script>
        // ============================================
        // MAIN APPLICATION LOGIC
        // ============================================
        
        let selectedFiles = [];
        let parsedData = [];
        let startTime;
        let currentFilter = 'all';
        
        const fileInput = document.getElementById('fileInput');
        const uploadArea = document.getElementById('uploadArea');
        const fileList = document.getElementById('fileList');
        const parseBtn = document.getElementById('parseBtn');
        const clearBtn = document.getElementById('clearBtn');
        const progressContainer = document.getElementById('progressContainer');
        const progressBar = document.getElementById('progressBar');
        const progressText = document.getElementById('progressText');
        const successAlert = document.getElementById('successAlert');
        const errorAlert = document.getElementById('errorAlert');
        const warningAlert = document.getElementById('warningAlert');
        const resultsContainer = document.getElementById('resultsContainer');
        
        {% if not default_theme %}
        // Enhanced theme toggle with smooth transition
        function toggleTheme() {
            const body = document.body;
            const currentTheme = body.getAttribute('data-theme');
            const newTheme = currentTheme === 'light' ? 'dark' : 'light';
            
            body.style.transition = 'background-color 0.4s ease, color 0.4s ease';
            body.setAttribute('data-theme', newTheme);
            
            const themeToggle = document.querySelector('.theme-toggle');
            if (themeToggle) {
                themeToggle.textContent = newTheme === 'light' ? '🌙 Dark' : '☀️ Light';
                themeToggle.style.animation = 'rotate 0.6s ease';
            }
            
            localStorage.setItem('theme', newTheme);
        }

        // Load saved theme
        window.addEventListener('DOMContentLoaded', () => {
            const themeToggle = document.querySelector('.theme-toggle');
            const savedTheme = localStorage.getItem('theme') || 'light';
            document.body.setAttribute('data-theme', savedTheme);
            if (themeToggle) themeToggle.textContent = savedTheme === 'light' ? '🌙' : '☀️';
        });
        {% else %}
        // Server-enforced theme: client-side toggle omitted to reduce payload
        {% endif %}
        
        const MAX_UPLOADS = {{max_uploads}};
        const MAX_FILE_SIZE_MB = {{max_file_size_mb}};

        // File input change
        fileInput.addEventListener('change', function(e) {
            handleFiles(e.target.files);
        });
        
        // Drag and drop
        uploadArea.addEventListener('dragover', function(e) {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });
        
        uploadArea.addEventListener('dragleave', function(e) {
            uploadArea.classList.remove('dragover');
        });
        
        uploadArea.addEventListener('drop', function(e) {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            handleFiles(e.dataTransfer.files);
        });
        
        function handleFiles(files) {
            const all = Array.from(files);
            const allowed = [];
            const skipped = [];
            const maxBytes = MAX_FILE_SIZE_MB * 1024 * 1024;
            for (const f of all) {
                if (f.size > maxBytes) {
                    skipped.push(f.name);
                } else {
                    allowed.push(f);
                }
            }
            if (skipped.length > 0) {
                showWarning('Some files were skipped because they exceed the max size: ' + skipped.join(', '));
            }
            selectedFiles = allowed.slice(0, MAX_UPLOADS);
            // Render the file list, hide the bottom parse button and auto-start parsing
            if (selectedFiles.length > 0) {
                clearBtn.style.display = 'block';
                parseBtn.style.display = 'none'; // remove manual parse option
            } else {
                parseBtn.style.display = 'none';
                clearBtn.style.display = 'none';
            }

            updateStats();

            hideAlerts();
            resultsContainer.style.display = 'none';

            displayFileList();

            // Auto-start parsing after a short delay to allow UI render
            if (selectedFiles.length > 0) {
                setTimeout(() => { parseResumes(); }, 300);
            }
        }

        function displayFileList() {
            let html = '<div class="file-list-header">';
            html += '<div class="file-list-title">📁 Selected Files (' + selectedFiles.length + ')</div>';
            html += '<button class="clear-files" onclick="clearFiles()">Clear All</button>';
            html += '</div>';

            selectedFiles.forEach((file, index) => {
                const sizeMB = (file.size / 1024 / 1024).toFixed(2);
                html += '<div class="file-item" id="file-' + index + '">';
                html += '<span class="file-icon">📄</span>';
                html += '<div class="file-info" id="info-' + index + '">';
                html += '<div class="file-name">' + file.name + '</div>';
                html += '<div class="file-size">' + sizeMB + ' MB</div>';
                html += '</div>';
                html += '<div class="file-status pending" id="status-' + index + '">⏳ Pending</div>';
                html += '<div class="file-progress" id="progress-' + index + '" style="width:0%"></div>';
                html += '</div>';
            });

            fileList.innerHTML = html;
            fileList.style.display = 'block';
        }

        function clearFiles() {
            selectedFiles = [];
            fileInput.value = '';
            parseBtn.disabled = true;
            clearBtn.style.display = 'none';
            // keep file-list hidden
            fileList.style.display = 'none';
            updateStats();
            hideAlerts();
        }

        function clearAll() {
            clearFiles();
            parsedData = [];
            resultsContainer.style.display = 'none';
            progressContainer.style.display = 'none';
            updateStats();
        }

        async function clearParsed() {
            // remove server-side preview text files (does not touch originals)
            try {
                const resp = await fetch('/clear_parsed', { method: 'POST' });
                if (!resp.ok) throw new Error('Failed to clear parsed files');
                const body = await resp.json();
                parsedData = [];
                resultsContainer.style.display = 'none';
                fileList.style.display = 'none';
                clearFiles();
                showSuccess('🧹 Cleared parsed previews (' + (body.deleted || 0) + ')');
            } catch (err) {
                showError('Error clearing parsed files: ' + err.message);
            }
        }

        function updateStats() {
            document.getElementById('totalFiles').textContent = selectedFiles.length;
            
            const successful = parsedData.filter(d => !d.full_name?.startsWith('Error:')).length;
            const failed = parsedData.filter(d => d.full_name?.startsWith('Error:')).length;
            
            document.getElementById('successCount').textContent = successful;
            document.getElementById('failedCount').textContent = failed;
        }
        
        async function parseResumes() {
            if (selectedFiles.length === 0) {
                showError('Please select at least one resume file');
                return;
            }
            
            hideAlerts();
            resultsContainer.style.display = 'none';
            progressContainer.style.display = 'block';
            parseBtn.disabled = true;
            clearBtn.disabled = true;
            
            parsedData = [];
            let successCount = 0;
            let failedCount = 0;
            startTime = Date.now();

            // Use a single progress display; send one file per request sequentially
            for (let i = 0; i < selectedFiles.length; i++) {
                const file = selectedFiles[i];
                const formData = new FormData();
                formData.append('files', file);

                // per-file UI elements
                const statusEl = document.getElementById('status-' + i);
                const infoEl = document.getElementById('info-' + i);
                const progressEl = document.getElementById('progress-' + i);

                if (statusEl) {
                    statusEl.className = 'file-status parsing';
                    statusEl.innerHTML = '⚡ Parsing...';
                }
                if (progressEl) progressEl.style.width = '50%';

                try {
                    // include optional per-request API key and model
                    const headers = {};
                    const storedKey = sessionStorage.getItem('GROK_API_KEY');
                    if (document.getElementById('llmToggle').checked && storedKey) headers['X-API-KEY'] = storedKey;
                    headers['X-MODEL'] = document.getElementById('modelSelect').value || 'gpt-4o-mini';

                    const response = await fetch('/parse', { method: 'POST', body: formData, headers: headers });
                    if (!response.ok) throw new Error('Failed to parse');

                    const data = await response.json();
                    const result = data.results[0];
                    parsedData.push(result);

                    // replace file info with parsed data
                    if (infoEl) {
                        if (result.full_name?.startsWith('Error:')) {
                            infoEl.innerHTML = '<div class="file-name">Parsing failed</div><div class="small">' + (result.full_name || '') + '</div>';
                        } else {
                            // Only show the candidate name and a link to view the uploaded file (no inline preview)
                            let html = '<div class="file-name">' + (result.full_name || '-') + '</div>';
                            if (result.original_link) {
                                html += '<div class="small"><a href="' + result.original_link + '" target="_blank">View original</a></div>';
                            } else if (result.file_link) {
                                html += '<div class="small"><a href="' + result.file_link + '" target="_blank">View file</a></div>';
                            }
                            infoEl.innerHTML = html;
                        }
                    }

                    if (result.full_name?.startsWith('Error:')) {
                        failedCount++;
                        if (statusEl) { statusEl.className = 'file-status failed'; statusEl.innerHTML = '❌ Failed'; }
                    } else {
                        successCount++;
                        if (statusEl) { statusEl.className = 'file-status success'; statusEl.innerHTML = '✅ Parsed'; }
                    }

                    if (progressEl) progressEl.style.width = '100%';

                } catch (error) {
                    parsedData.push({
                        full_name: 'Error: ' + file.name,
                        email: null,
                        phone_number: null,
                        alternate_phone_number: null,
                        highest_qualification: null,
                        years_of_experience: null,
                        current_company: null,
                        current_designation: null,
                        city: null,
                        state: null
                    });
                    failedCount++;
                    if (infoEl) infoEl.innerHTML = '<div class="file-name">Parsing error</div><div class="small">' + error.message + '</div>';
                    if (statusEl) { statusEl.className = 'file-status failed'; statusEl.innerHTML = '❌ Failed'; }
                    if (progressEl) progressEl.style.width = '100%';
                }

                // Update aggregate progress
                const progress = ((i + 1) / selectedFiles.length) * 100;
                progressBar.style.width = progress + '%';
                progressBar.textContent = Math.round(progress) + '%';
                progressText.textContent = (i + 1) + '/' + selectedFiles.length;

                // Update counts and elapsed time
                updateStats();
                const elapsed = ((Date.now() - startTime) / 1000).toFixed(1);
                document.getElementById('processingTime').textContent = elapsed + 's';
            }
            
            // Complete
            progressBar.classList.add('complete');
            parseBtn.disabled = false;
            clearBtn.disabled = false;
            
            const totalTime = ((Date.now() - startTime) / 1000).toFixed(1);
            document.getElementById('processingTime').textContent = totalTime + 's';
            
            // Show results
            if (successCount > 0) {
                showSuccess('✅ Successfully parsed ' + successCount + ' resume(s)!');
            }
            
            if (failedCount > 0) {
                showWarning('⚠️ ' + failedCount + ' resume(s) failed to parse. Check the file list for details.');
            }
            
            if (parsedData.length > 0) {
                displayResults(parsedData);
            }
        }
        
        function displayResults(results) {
            resultsContainer.style.display = 'block';
            currentFilter = 'all';
            updateFilterTabs();
            renderTable(results);
        }

        function renderTable(results) {
            let html = '<thead><tr>';
            html += '<th>#</th>';
            html += '<th>Status</th>';
            html += '<th>Full Name</th>';
            html += '<th>Email</th>';
            html += '<th>Phone Number</th>';
            html += '<th>Alt. Phone</th>';
            html += '<th>Qualification</th>';
            html += '<th>Experience (Yrs)</th>';
            html += '<th>Current Company</th>';
            html += '<th>Designation</th>';
            html += '<th>City</th>';
                html += '<th>State</th>';
            
            html += '</tr></thead><tbody>';
            
            results.forEach((row, idx) => {
                const isFailed = row.full_name?.startsWith('Error:');
                const isIncomplete = !isFailed && countNullFields(row) > 5;
                
                html += '<tr>';
                html += '<td>' + (idx + 1) + '</td>';
                
                if (isFailed) {
                    html += '<td><span style="color: var(--danger);">❌ Failed</span></td>';
                } else if (isIncomplete) {
                    html += '<td><span style="color: var(--warning);">⚠️ Incomplete</span></td>';
                } else {
                    html += '<td><span style="color: var(--success);">✅ Success</span></td>';
                }
                
                let nameCell = formatValue(row.full_name);
                if (row.original_link) {
                    nameCell = '<a href="' + row.original_link + '" target="_blank">' + (row.full_name || '-') + '</a>';
                } else if (row.file_link) {
                    nameCell = '<a href="' + row.file_link + '" target="_blank">' + (row.full_name || '-') + '</a>';
                }
                html += '<td>' + nameCell + '</td>';
                html += '<td>' + formatValue(row.email) + '</td>';
                html += '<td>' + formatValue(row.phone_number) + '</td>';
                html += '<td>' + formatValue(row.alternate_phone_number) + '</td>';
                html += '<td>' + formatValue(row.highest_qualification) + '</td>';
                html += '<td>' + formatValue(row.years_of_experience) + '</td>';
                html += '<td>' + formatValue(row.current_company) + '</td>';
                html += '<td>' + formatValue(row.current_designation) + '</td>';
                html += '<td>' + formatValue(row.city) + '</td>';
                html += '<td>' + formatValue(row.state) + '</td>';
                
                        html += '</tr>';
            });
            
            html += '</tbody>';
            document.getElementById('resultsTable').innerHTML = html;
        }

        

        function countNullFields(row) {
            let count = 0;
            for (let key in row) {
                if (row[key] === null || row[key] === undefined || row[key] === '') {
                    count++;
                }
            }
            return count;
        }
        
        function formatValue(value) {
            if (value === null || value === undefined || value === '') {
                return '<span class="null-value">-</span>';
            }
            return value;
        }

        function filterResults(type) {
            currentFilter = type;
            updateFilterTabs();
            
            let filtered = parsedData;
            
            if (type === 'success') {
                filtered = parsedData.filter(d => !d.full_name?.startsWith('Error:') && countNullFields(d) <= 5);
            } else if (type === 'failed') {
                filtered = parsedData.filter(d => d.full_name?.startsWith('Error:'));
            } else if (type === 'incomplete') {
                filtered = parsedData.filter(d => !d.full_name?.startsWith('Error:') && countNullFields(d) > 5);
            }
            
            renderTable(filtered);
        }

        function updateFilterTabs() {
            const tabs = document.querySelectorAll('.filter-tab');
            tabs.forEach((tab, idx) => {
                tab.classList.remove('active');
                const filters = ['all', 'success', 'failed', 'incomplete'];
                if (filters[idx] === currentFilter) {
                    tab.classList.add('active');
                }
            });
        }

        function searchTable() {
            const input = document.getElementById('searchBox').value.toLowerCase();
            const table = document.getElementById('resultsTable');
            const rows = table.getElementsByTagName('tr');
            
            for (let i = 1; i < rows.length; i++) {
                const row = rows[i];
                const text = row.textContent.toLowerCase();
                
                if (text.includes(input)) {
                    row.style.display = '';
                } else {
                    row.style.display = 'none';
                }
            }
        }
        
        async function downloadExcel() {
            if (parsedData.length === 0) return;
            
            try {
                const response = await fetch('/export', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({data: parsedData})
                });
                
                if (!response.ok) {
                    throw new Error('Failed to export to Excel');
                }
                
                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = 'parsed_resumes_' + new Date().getTime() + '.xlsx';
                document.body.appendChild(a);
                a.click();
                document.body.removeChild(a);
                window.URL.revokeObjectURL(url);
                
                showSuccess('📥 Excel file downloaded successfully!');
                
            } catch (error) {
                showError('Error downloading Excel: ' + error.message);
            }
        }

        async function downloadJSON() {
            if (parsedData.length === 0) return;
            
            try {
                const dataStr = JSON.stringify(parsedData, null, 2);
                const blob = new Blob([dataStr], { type: 'application/json' });
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = 'parsed_resumes_' + new Date().getTime() + '.json';
                document.body.appendChild(a);
                a.click();
                document.body.removeChild(a);
                window.URL.revokeObjectURL(url);
                
                showSuccess('📥 JSON file downloaded successfully!');
                
            } catch (error) {
                showError('Error downloading JSON: ' + error.message);
            }
        }
        
        function showSuccess(message) {
            successAlert.textContent = message;
            successAlert.style.display = 'block';
            errorAlert.style.display = 'none';
            warningAlert.style.display = 'none';
            setTimeout(() => { successAlert.style.display = 'none'; }, 5000);
        }
        
        function showError(message) {
            errorAlert.textContent = message;
            errorAlert.style.display = 'block';
            successAlert.style.display = 'none';
            warningAlert.style.display = 'none';
        }

        function showWarning(message) {
            warningAlert.textContent = message;
            warningAlert.style.display = 'block';
            successAlert.style.display = 'none';
            errorAlert.style.display = 'none';
        }
        
        function hideAlerts() {
            successAlert.style.display = 'none';
            errorAlert.style.display = 'none';
            warningAlert.style.display = 'none';
        }

        // Settings modal logic
        function openSettings() {
            document.getElementById('modalBackdrop').style.display = 'flex';
            // load saved settings
            const savedModel = localStorage.getItem('model') || 'gpt-4o-mini';
            document.getElementById('modelSelect').value = savedModel;
            const snowVal = localStorage.getItem('snow-intensity') || '120';
            document.getElementById('snowIntensity').value = snowVal;
            const llm = localStorage.getItem('use_llm') === '1';
            document.getElementById('llmToggle').checked = llm;
        }

        function closeSettings() {
            document.getElementById('modalBackdrop').style.display = 'none';
        }

        function saveSettings() {
            const model = document.getElementById('modelSelect').value;
            const snowVal = document.getElementById('snowIntensity').value;
            const use_llm = document.getElementById('llmToggle').checked;
            const key = document.getElementById('apiKeyInput').value;
            localStorage.setItem('model', model);
            localStorage.setItem('snow-intensity', snowVal);
            localStorage.setItem('use_llm', use_llm ? '1' : '0');
            if (use_llm && key) {
                // store in session for server requests
                sessionStorage.setItem('GROK_API_KEY', key);
            }
            // apply snow intensity
            if (snowEnabled) {
                stopSnow(); startSnow();
            }
            closeSettings();
            showSuccess('Settings saved');
        }

        function clearLocalSettings() {
            localStorage.removeItem('model');
            localStorage.removeItem('snow-intensity');
            localStorage.removeItem('use_llm');
            sessionStorage.removeItem('GROK_API_KEY');
            showSuccess('Settings reset');
        }

        

        function downloadSample() {
            const sample = `John Doe\njohn.doe@example.com\n+1 (555) 123-4567\n\nExperience\nSenior Developer at Acme Corp, Jan 2020 - Present\n\nEducation\nMaster of Science in Computer Science\n\n6 years of experience\n\nSpringfield, Illinois`;
            const blob = new Blob([sample], {type:'text/plain'});
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a'); a.href = url; a.download = 'sample_resume.txt'; a.click(); URL.revokeObjectURL(url);
        }

        // Preview modal
        function showPreview(obj) {
            document.getElementById('previewContent').textContent = JSON.stringify(obj, null, 2);
            document.getElementById('previewBackdrop').style.display = 'flex';
        }

        function closePreview() {
            document.getElementById('previewBackdrop').style.display = 'none';
        }
        
        // --- Snow animation ---
        let snowEnabled = true;
        let snowCanvas = null;
        let ctx = null;
        let flakes = [];
        let snowAnimId = null;

        function initSnowCanvas() {
            snowCanvas = document.getElementById('snowCanvas');
            if (!snowCanvas) {
                snowCanvas = document.createElement('canvas');
                snowCanvas.id = 'snowCanvas';
                document.body.insertBefore(snowCanvas, document.body.firstChild);
            }
            ctx = snowCanvas.getContext ? snowCanvas.getContext('2d') : null;
            resizeSnow();
        }

        function resizeSnow() {
            if (!snowCanvas) return;
            snowCanvas.width = window.innerWidth;
            snowCanvas.height = window.innerHeight;
        }

        function createFlakes(count) {
            flakes = [];
            for (let i = 0; i < count; i++) {
                flakes.push({
                    x: Math.random() * snowCanvas.width,
                    y: Math.random() * snowCanvas.height,
                    r: 1 + Math.random() * 4,
                    d: 0.5 + Math.random() * 1.5
                });
            }
        }

        function drawSnow() {
            if (!ctx || !snowEnabled) return;
            ctx.clearRect(0, 0, snowCanvas.width, snowCanvas.height);
            for (let i = 0; i < flakes.length; i++) {
                const f = flakes[i];
                ctx.beginPath();
                ctx.fillStyle = 'rgba(255,255,255,' + (0.6 + Math.random()*0.35) + ')';
                ctx.arc(f.x, f.y, f.r, 0, Math.PI * 2, true);
                ctx.fill();
            }
            updateFlakes();
        }

        function updateFlakes() {
            for (let i = 0; i < flakes.length; i++) {
                const f = flakes[i];
                f.y += Math.pow(f.d, 1.2) + 0.3 + f.r * 0.08;
                f.x += Math.sin(f.y * 0.01) * (0.5 + f.r * 0.02);

                if (f.y > snowCanvas.height + 10) {
                    flakes[i] = { x: Math.random() * snowCanvas.width, y: -10 - Math.random()*50, r: f.r, d: f.d };
                }
            }
        }

        function animateSnow() {
            drawSnow();
            snowAnimId = requestAnimationFrame(animateSnow);
        }

        function startSnow() {
            if (!snowCanvas || !ctx) initSnowCanvas();
            if (!ctx) return;
            resizeSnow();
            createFlakes(Math.max(120, Math.floor(window.innerWidth / 12)));
            if (!snowAnimId) animateSnow();
        }

        function stopSnow() {
            if (snowAnimId) cancelAnimationFrame(snowAnimId);
            snowAnimId = null;
            if (ctx && snowCanvas) ctx.clearRect(0, 0, snowCanvas.width, snowCanvas.height);
        }

        function toggleSnow() {
            snowEnabled = !snowEnabled;
            const btn = document.getElementById('snowToggle');
            if (btn) btn.textContent = snowEnabled ? '❄️ Snow' : '🌨️ Off';
            if (snowEnabled) startSnow(); else stopSnow();
            localStorage.setItem('snow', snowEnabled ? '1' : '0');
        }

        // expose controls for console debugging
        window.startSnow = startSnow;
        window.stopSnow = stopSnow;
        window.toggleSnow = toggleSnow;

        window.addEventListener('resize', function() { if (snowEnabled) resizeSnow(); });
        window.addEventListener('DOMContentLoaded', function() {
            const saved = localStorage.getItem('snow');
            if (saved === '0') {
                snowEnabled = false;
                const btn = document.getElementById('snowToggle'); if (btn) btn.textContent = '🌨️ Off';
            } else {
                snowEnabled = true;
            }
            if (snowEnabled) startSnow();
        });
    </script>
</body>
</html>
'''

@app.route('/')
def index():
    max_uploads = int(os.getenv('PARSE_MAX_UPLOADS', '500'))
    max_file_size_mb = int(os.getenv('PARSE_MAX_FILE_MB', '5'))
    # Allow enforcing a default theme via DEFAULT_THEME env var ('dark' or 'light')
    raw_theme = os.getenv('DEFAULT_THEME', '').lower()
    default_theme = raw_theme if raw_theme in ('dark', 'light') else ''
    return render_template_string(
        HTML_TEMPLATE,
        max_uploads=max_uploads,
        max_file_size_mb=max_file_size_mb,
        default_theme=default_theme
    )

@app.route('/parse', methods=['POST'])
def parse():
    """Parse resume files with security validation"""
    try:
        files = request.files.getlist('files')
        if not files:
            return jsonify({'error': 'No files uploaded'}), 400
        
        # Validate file count
        max_uploads = int(os.getenv('PARSE_MAX_UPLOADS', '500'))
        if len(files) > max_uploads:
            return jsonify({'error': f'Too many files (max: {max_uploads})'}), 400
        
        # Validate each file
        valid_extensions = {'.txt', '.pdf', '.doc', '.docx'}
        max_size = int(os.getenv('PARSE_MAX_FILE_MB', '5')) * 1024 * 1024
        
        for file in files:
            # Validate filename
            if not file.filename:
                return jsonify({'error': 'Invalid filename'}), 400
            
            # Check extension
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in valid_extensions:
                return jsonify({'error': f'Invalid file type: {ext}. Allowed: {", ".join(valid_extensions)}'}), 400
            
            # Check file size (in-memory)
            file.seek(0, os.SEEK_END)
            file_size = file.tell()
            file.seek(0)
            if file_size > max_size:
                return jsonify({'error': f'File too large (max: {max_size / (1024*1024):.0f} MB)'}), 400
            elif file_size == 0:
                return jsonify({'error': 'Empty file uploaded'}), 400
        
        # per-request API key and model (from headers) - sanitize
        req_api_key = request.headers.get('X-API-KEY', '').strip()[:200]  # Truncate to prevent abuse
        req_model = (request.headers.get('X-MODEL') or request.form.get('model') or '').strip()[:100]
        
        # Validate model to prevent injection
        allowed_models = ['gpt-4o-mini', 'gpt-4', 'grok-1']
        if req_model and req_model not in allowed_models:
            req_model = 'gpt-4o-mini'

        uploads_dir = os.path.join(os.getcwd(), 'uploads')
        originals_dir = os.path.join(uploads_dir, 'originals')
        
        # Verify directories are within cwd (path traversal protection)
        try:
            uploads_dir = os.path.realpath(uploads_dir)
            originals_dir = os.path.realpath(originals_dir)
            if not uploads_dir.startswith(os.path.realpath(os.getcwd())):
                raise ValueError('Invalid uploads directory')
            if not originals_dir.startswith(os.path.realpath(os.getcwd())):
                raise ValueError('Invalid originals directory')
        except (ValueError, OSError) as e:
            return jsonify({'error': 'Invalid directory configuration'}), 500
        
        os.makedirs(uploads_dir, exist_ok=True)
        os.makedirs(originals_dir, exist_ok=True)

        # batching and concurrency settings
        batch_size = int(os.getenv('PARSE_BATCH_SIZE', '50'))
        workers = int(os.getenv('PARSE_WORKERS', '4'))
        llm_concurrency = int(os.getenv('PARSE_LLM_CONCURRENCY', '2'))
        save_originals = os.getenv('STORE_ORIGINALS', '0') == '1'

        llm_sema = threading.Semaphore(llm_concurrency)

        results = [None] * len(files)

        def _sanitize_filename(filename):
            """Safely sanitize filename to prevent path traversal and injection"""
            # Remove any path components
            filename = os.path.basename(filename)
            # Keep only safe characters
            safe_name = re.sub(r'[^A-Za-z0-9._-]', '_', filename)
            # Limit length
            safe_name = safe_name[:150]
            # Prevent empty names
            if not safe_name or safe_name.startswith('.'):
                safe_name = 'resume_' + datetime.now().strftime('%Y%m%d%H%M%S%f')
            return safe_name

        def _process(idx, file_storage):
            try:
                # optionally save original
                orig_name = None
                if save_originals:
                    try:
                        safe_name = _sanitize_filename(file_storage.filename)
                        ts = datetime.now().strftime('%Y%m%d%H%M%S%f')
                        orig_name = f"{ts}_{safe_name}"
                        orig_path = os.path.join(originals_dir, orig_name)
                        
                        # Verify path is still within originals_dir
                        orig_path_real = os.path.realpath(orig_path)
                        if not orig_path_real.startswith(os.path.realpath(originals_dir)):
                            raise ValueError('Path traversal attempt detected')
                        
                        file_storage.stream.seek(0)
                        file_storage.save(orig_path_real)
                    except Exception:
                        orig_name = None

                try:
                    file_storage.stream.seek(0)
                except Exception:
                    pass
                content = read_file_content(file_storage)

                # Save a reduced text preview
                text_fname = None
                try:
                    safe_name = _sanitize_filename(file_storage.filename)
                    ts = datetime.now().strftime('%Y%m%d%H%M%S%f')
                    text_fname = f"{ts}_{safe_name}.txt"
                    text_path = os.path.join(uploads_dir, text_fname)
                    
                    # Verify path is still within uploads_dir
                    text_path_real = os.path.realpath(text_path)
                    if not text_path_real.startswith(os.path.realpath(uploads_dir)):
                        raise ValueError('Path traversal attempt detected')
                    
                    # Limit content size to prevent disk exhaustion
                    max_content = 1024 * 500  # 500 KB
                    with open(text_path_real, 'w', encoding='utf-8') as tf:
                        tf.write(content[:max_content])
                except Exception:
                    text_fname = None

                parsed = parse_resume_text(content)

                llm_result = None
                if call_llm_extract is not None:
                    try:
                        acquired = llm_sema.acquire(timeout=30)
                        if acquired:
                            try:
                                llm_out = call_llm_extract(content, api_key=req_api_key, model=req_model or 'gpt-4o-mini')
                                if llm_out and isinstance(llm_out, list) and len(llm_out) > 0:
                                    llm_result = llm_out[0]
                            finally:
                                llm_sema.release()
                    except Exception:
                        llm_result = None

                if llm_result and isinstance(llm_result, dict):
                    merged = {}
                    fields = [
                        'full_name', 'email', 'phone_number', 'alternate_phone_number',
                        'highest_qualification', 'years_of_experience', 'current_company',
                        'current_designation', 'city', 'state'
                    ]
                    for f in fields:
                        local_val = parsed.get(f)
                        llm_field = llm_result.get(f)
                        chosen = local_val
                        if isinstance(llm_field, dict):
                            conf = float(llm_field.get('confidence') or 0.0)
                            if conf >= float(os.getenv('CONF_THRESHOLD', '0.8')) and llm_field.get('value') is not None:
                                chosen = llm_field.get('value')
                        elif llm_field is not None:
                            chosen = llm_field
                        merged[f] = chosen
                    if text_fname:
                        merged['file_link'] = '/uploads/' + text_fname
                    if orig_name:
                        merged['original_link'] = '/originals/' + orig_name
                    results[idx] = merged
                else:
                    if text_fname:
                        parsed['file_link'] = '/uploads/' + text_fname
                    if orig_name:
                        parsed['original_link'] = '/originals/' + orig_name
                    results[idx] = parsed

            except Exception:
                results[idx] = {
                    'full_name': f'Error: {file_storage.filename}',
                    'email': None,
                    'phone_number': None,
                    'alternate_phone_number': None,
                    'highest_qualification': None,
                    'years_of_experience': None,
                    'current_company': None,
                    'current_designation': None,
                    'city': None,
                    'state': None
                }

        total = len(files)
        for start in range(0, total, batch_size):
            batch = files[start:start+batch_size]
            with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
                futures = []
                for idx_offset, file in enumerate(batch):
                    idx = start + idx_offset
                    futures.append(executor.submit(_process, idx, file))
                concurrent.futures.wait(futures)

        return jsonify({'results': results})
    
    except Exception as e:
        # Don't leak detailed error messages in production
        import logging
        logging.error(f'Parse error: {str(e)}', exc_info=True)
        
        if os.getenv('FLASK_ENV') == 'development':
            return jsonify({'error': str(e)}), 500
        else:
            return jsonify({'error': 'An error occurred while processing files'}), 500


def _check_admin_token() -> bool:
    """Simple admin auth using ADMIN_TOKEN env var and X-ADMIN-TOKEN header."""
    admin_token = os.getenv('ADMIN_TOKEN')
    if not admin_token:
        return False
    header = request.headers.get('X-ADMIN-TOKEN') or request.args.get('admin_token')
    return header == admin_token


@app.route('/admin/set_api_key', methods=['POST'])
def admin_set_api_key():
    if not _check_admin_token():
        return jsonify({'error': 'unauthorized'}), 401

    if store_api_key is None:
        return jsonify({'error': 'server-side storage is not configured (MASTER_KEY missing)'}), 500

    data = request.json or {}
    api_key = data.get('api_key')
    if not api_key:
        return jsonify({'error': 'api_key required in JSON body'}), 400

    try:
        store_api_key(api_key)
        return jsonify({'status': 'ok'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/admin/has_api_key', methods=['GET'])
def admin_has_api_key():
    if not _check_admin_token():
        return jsonify({'error': 'unauthorized'}), 401

    try:
        exists = False
        if get_stored_api_key is not None:
            exists = get_stored_api_key() is not None
        return jsonify({'stored': exists})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/admin/delete_api_key', methods=['POST', 'DELETE'])
def admin_delete_api_key():
    if not _check_admin_token():
        return jsonify({'error': 'unauthorized'}), 401

    if delete_api_key is None:
        return jsonify({'error': 'server-side storage is not configured (MASTER_KEY missing)'}), 500

    try:
        delete_api_key()
        return jsonify({'status': 'deleted'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/export', methods=['POST'])
def export():
    try:
        data = request.json.get('data', [])
        
        if not data:
            return jsonify({'error': 'No data to export'}), 400
        
        # Create DataFrame with exact column order
        columns = [
            'full_name',
            'email',
            'phone_number',
            'alternate_phone_number',
            'highest_qualification',
            'years_of_experience',
            'current_company',
            'current_designation',
            'city',
            'state'
        ]
        
        df = pd.DataFrame(data, columns=columns)
        
        # Create Excel file in memory
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Parsed Resumes')
            
            # Auto-adjust column widths
            worksheet = writer.sheets['Parsed Resumes']
            for idx, col in enumerate(df.columns):
                max_length = max(
                    df[col].astype(str).map(len).max(),
                    len(col)
                ) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = min(max_length, 50)
        
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'parsed_resumes_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/clear_parsed', methods=['POST'])
def clear_parsed():
    """Delete server-side parsed preview files in `uploads/` but leave `uploads/originals/` intact."""
    uploads_dir = os.path.join(os.getcwd(), 'uploads')
    deleted = 0
    try:
        if os.path.exists(uploads_dir):
            for fname in os.listdir(uploads_dir):
                fpath = os.path.join(uploads_dir, fname)
                # skip originals directory
                if os.path.isdir(fpath) and os.path.basename(fpath) == 'originals':
                    continue
                try:
                    if os.path.isfile(fpath):
                        os.remove(fpath)
                        deleted += 1
                except Exception:
                    pass
        return jsonify({'status': 'ok', 'deleted': deleted})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


 
def read_file_content(file):
    """
    Read content from uploaded file with robust error handling.
    Supports: TXT, PDF, DOCX, DOC with multiple fallback methods.
    """
    try:
        if not file or not file.filename:
            return None
        
        filename = file.filename.lower()
        file.seek(0)  # Reset file pointer
        
        # TXT files
        if filename.endswith('.txt'):
            try:
                content = file.read()
                text = content.decode('utf-8', errors='replace')
                return text if text.strip() else None
            except Exception as e:
                logging.warning(f"Failed to read TXT: {e}")
                return None
        
        # PDF files - multiple methods
        if filename.endswith('.pdf'):
            # Method 1: PyPDF2
            try:
                import PyPDF2
                file.seek(0)
                pdf_reader = PyPDF2.PdfReader(file)
                if pdf_reader.pages:
                    pages = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            text = page.extract_text()
                            if text and text.strip():
                                pages.append(text)
                        except Exception:
                            logging.warning(f"Failed to extract page {page_num}")
                    if pages:
                        result = '\n'.join(pages)
                        return result if result.strip() else None
            except Exception as e:
                logging.warning(f"PyPDF2 failed: {e}")
            
            # Method 2: pdfplumber
            try:
                import pdfplumber
                file.seek(0)
                with pdfplumber.open(file) as pdf:
                    pages = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            pages.append(text)
                    if pages:
                        result = '\n'.join(pages)
                        return result if result.strip() else None
            except Exception:
                logging.debug("pdfplumber not available")
            
            # Fallback: binary read
            try:
                file.seek(0)
                content = file.read()
                text = content.decode('utf-8', errors='replace')
                return text if text.strip() else None
            except Exception as e:
                logging.warning(f"PDF binary read failed: {e}")
        
        # DOCX files
        if filename.endswith('.docx'):
            try:
                file.seek(0)
                text = docx2txt.process(file)
                return text if text and text.strip() else None
            except Exception as e:
                logging.warning(f"docx2txt failed: {e}")
            
            # Fallback: python-docx
            try:
                from docx import Document
                from io import BytesIO
                file.seek(0)
                doc = Document(BytesIO(file.read()))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                if paragraphs:
                    result = '\n'.join(paragraphs)
                    return result if result.strip() else None
            except Exception as e:
                logging.warning(f"python-docx failed: {e}")
        
        # DOC files (older Word format)
        if filename.endswith('.doc'):
            try:
                from docx import Document
                from io import BytesIO
                file.seek(0)
                doc = Document(BytesIO(file.read()))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                if paragraphs:
                    result = '\n'.join(paragraphs)
                    return result if result.strip() else None
            except Exception as e:
                logging.warning(f"DOC read failed: {e}")
        
        # Fallback: try binary decode
        try:
            file.seek(0)
            content = file.read()
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = content.decode(encoding, errors='replace')
                    if text.strip():
                        return text
                except Exception:
                    continue
        except Exception as e:
            logging.error(f"Fallback decode failed: {e}")
        
        return None
    
    except Exception as e:
        logging.error(f"Unexpected error in read_file_content: {e}")
        return None


@app.route('/uploads/<path:fname>')
def uploaded_file(fname):
    """Serve uploaded files with path traversal protection"""
    uploads_dir = os.path.realpath(os.path.join(os.getcwd(), 'uploads'))
    
    try:
        # Prevent path traversal attacks
        safe_path = os.path.realpath(os.path.join(uploads_dir, fname))
        
        if not safe_path.startswith(uploads_dir):
            return abort(403)
        
        if not os.path.exists(safe_path) or not os.path.isfile(safe_path):
            return abort(404)
        
        return send_file(safe_path, as_attachment=False)
    except Exception:
        return abort(400)


@app.route('/originals/<path:fname>')
def original_file(fname):
    """Serve original files with path traversal protection"""
    originals_dir = os.path.realpath(os.path.join(os.getcwd(), 'uploads', 'originals'))
    
    try:
        # Prevent path traversal attacks
        safe_path = os.path.realpath(os.path.join(originals_dir, fname))
        
        if not safe_path.startswith(originals_dir):
            return abort(403)
        
        if not os.path.exists(safe_path) or not os.path.isfile(safe_path):
            return abort(404)
        
        # Try to guess mimetype so browser can render inline when possible
        import mimetypes
        mime, _ = mimetypes.guess_type(safe_path)
        return send_file(safe_path, mimetype=mime or 'application/octet-stream', as_attachment=False)
    except Exception:
        return abort(400)

def parse_resume_text(text):
    """
    Parse resume text using the improved resume_parser module.
    """
    if parse_resume is None:
        # Fallback if import failed
        return {
            'full_name': None,
            'email': None,
            'phone_number': None,
            'alternate_phone_number': None,
            'highest_qualification': None,
            'years_of_experience': None,
            'current_company': None,
            'current_designation': None,
            'city': None,
            'state': None
        }
    
    # Use the improved resume_parser functions directly on text
    from resume_parser import (extract_name, extract_email, extract_phone,
                               extract_alternate_phone, extract_qualification,
                               extract_experience, extract_current_company,
                               extract_designation, extract_city, extract_state)
    
    result = {
        'full_name': extract_name(text),
        'email': extract_email(text),
        'phone_number': extract_phone(text),
        'alternate_phone_number': extract_alternate_phone(text),
        'highest_qualification': extract_qualification(text),
        'years_of_experience': extract_experience(text),
        'current_company': extract_current_company(text),
        'current_designation': extract_designation(text),
        'city': extract_city(text),
        'state': extract_state(text)
    }
    return result





def extract_email(text):
    """Extract email address - search entire document"""
    # Standard email pattern
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    
    if matches:
        # Validate and return first legitimate email
        for match in matches:
            email_clean = match.lower()
            # Skip obvious non-emails
            if '..' not in email_clean and not email_clean.startswith('.'):
                # Additional validation: reasonable length, not too many dots
                if len(email_clean) > 5 and email_clean.count('.') <= 3:
                    return email_clean
    
    return None

def extract_phone(text):
    """Extract primary phone number - handle diverse formats"""
    # Comprehensive patterns for different formats
    patterns = [
        r'\+\d{1,3}[\s.-]?\d{1,4}[\s.-]?\d{1,4}[\s.-]?\d{1,9}',           # +country code xxx xxx
        r'\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',                          # (xxx) xxx-xxxx or variants
        r'\b\d{10,11}\b',                                                  # 10-11 consecutive digits
        r'\d{3}[\s.-]\d{3}[\s.-]\d{4}',                                    # xxx-xxx-xxxx or xxx xxx xxxx
        r'\b\d{5}[\s.-]\d{5}\b',                                           # Indian: xxxxx xxxxx
    ]
    
    # Search entire document
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            phone = match.group(0)
            # Normalize: remove all non-digits except leading +
            normalized = re.sub(r'[^\d+]', '', phone)
            # Validate: must be 9-15 digits
            digit_count = len(re.sub(r'\D', '', normalized))
            if 9 <= digit_count <= 15:
                return normalized
    
    return None

def extract_alternate_phone(text):
    """Extract alternate phone number - look for second phone if exists"""
    phones = []
    
    # Same patterns as primary phone
    patterns = [
        r'\+\d{1,3}[\s.-]?\d{1,4}[\s.-]?\d{1,4}[\s.-]?\d{1,9}',
        r'\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',
        r'\b\d{10,11}\b',
        r'\d{3}[\s.-]\d{3}[\s.-]\d{4}',
        r'\b\d{5}[\s.-]\d{5}\b',
    ]
    
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            phone = match.group(0)
            normalized = re.sub(r'[^\d+]', '', phone)
            digit_count = len(re.sub(r'\D', '', normalized))
            if 9 <= digit_count <= 15 and normalized not in phones:
                phones.append(normalized)
    
    # Return second phone if exists, filter out duplicates
    unique_phones = []
    for p in phones:
        if p not in unique_phones:
            unique_phones.append(p)
    
    return unique_phones[1] if len(unique_phones) > 1 else None

def extract_qualification(text):
    """
    Extract highest qualification - ONLY return degree type, never institution name.
    Must find qualification in Education section for high confidence.
    """
    text_lower = text.lower()
    
    # Find Education section specifically
    edu_section = None
    lines = text.split('\n')
    
    edu_keywords = ['education', 'academic qualifications', 'qualifications', 'academic background', 'academic', 'schooling']
    
    for i, line in enumerate(lines):
        if any(kw in line.lower() for kw in edu_keywords):
            # Found education section, extract next 1500 chars
            edu_section = '\n'.join(lines[i:min(i+20, len(lines))])
            break
    
    if not edu_section:
        edu_section = text_lower
    
    # Qualification mappings (degree → keywords that identify it)
    qualifications = [
        ('PhD', ['phd', 'ph.d', 'ph d', 'doctorate', 'doctor of philosophy']),
        ('Masters', ['master', 'mba', 'ms', 'm.s', 'mtech', 'm.tech', 'msc', 'm.sc', 'postgraduate']),
        ('Bachelors', ['bachelor', 'b.tech', 'btech', 'b.e', 'be', 'bs', 'b.s', 'bsc', 'b.a', 'ba', 'bcom', 'b.com', 'undergraduate']),
        ('Diploma', ['diploma', 'dip', 'polytechnic']),
        ('Associate', ['associate', 'hsc', 'higher secondary'])
    ]
    
    # Search in education section ONLY
    for degree, keywords in qualifications:
        for kw in keywords:
            if kw in edu_section.lower():
                # Double-check: make sure we're not catching "Stanford Masters" 
                # by looking at context
                idx = edu_section.lower().find(kw)
                context = edu_section[max(0, idx-40):min(len(edu_section), idx+100)].lower()
                
                # Blocklist: if context contains university/college/school name indicators with "from/at/in"
                bad_patterns = [
                    r'from\s+[a-z\s]+university',
                    r'from\s+[a-z\s]+college',
                    r'from\s+[a-z\s]+school',
                    r'at\s+[a-z\s]+university',
                    r'in\s+[a-z\s]+university'
                ]
                
                found_bad = False
                for pattern in bad_patterns:
                    if re.search(pattern, context):
                        found_bad = True
                        break
                
                if not found_bad:
                    return degree
    
    return None

def extract_experience(text):
    """Extract years of experience - flexible patterns for different formats"""
    text_lower = text.lower()
    
    # Diverse patterns for experience detection
    patterns = [
        r'(\d+\.?\d*)\s*\+?\s*years?\s+(?:of\s+)?(?:work|professional|total)?\s*experience',
        r'experience\s*[:|=]\s*(\d+\.?\d*)\s*\+?\s*years?',
        r'(\d+\.?\d*)\s*\+?\s*(?:years?|yrs?|yrs?)\s+(?:work|professional)?(?:experience|exp)?',
        r'total\s+(?:work\s+)?experience\s*[:\-]?\s*(\d+\.?\d*)\s*years?',
        r'working\s+experience\s*[:\-]?\s*(\d+\.?\d*)\s*years?',
        r'years\s+(?:of\s+)?(?:work|professional|total)?\s*experience\s*[:\-]?\s*(\d+\.?\d*)',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            try:
                years = float(match.group(1))
                if 0 < years <= 80:  # Allow up to 80 years
                    return years
            except (ValueError, IndexError):
                continue
    
    return None

def extract_current_company(text):
    """Extract current company - search for multiple position indicators"""
    lines = text.split('\n')
    text_lower = text.lower()
    
    # Find experience section start
    exp_idx = -1
    for i, line in enumerate(lines):
        if 'experience' in line.lower() or 'work experience' in line.lower() or 'employment' in line.lower():
            exp_idx = i
            break
    
    if exp_idx == -1:
        # No experience section found, search entire document
        exp_idx = 0
    
    # Strategy 1: Look for "Present" keyword
    for j in range(exp_idx, min(exp_idx + 30, len(lines))):
        line = lines[j].strip()
        line_lower = line.lower()
        
        if 'present' in line_lower or 'current' in line_lower or 'till now' in line_lower:
            # Found current role, extract company name from nearby lines
            for offset in range(-3, 3):
                idx = j + offset
                if 0 <= idx < len(lines):
                    candidate = lines[idx].strip()
                    if candidate and 0 < len(candidate) < 100 and not candidate.isupper():
                        if len(candidate.split()) >= 1 and len(candidate.split()) <= 6:
                            # Check if looks like company name
                            if re.search(r'[A-Z][a-zA-Z]', candidate):
                                return candidate
    
    # Strategy 2: Look for company keywords
    company_indicators = ['company', 'organization', 'employer', 'firm', 'corporation', 'ltd', 'inc', 'corp']
    for j in range(exp_idx, min(exp_idx + 30, len(lines))):
        line = lines[j].strip()
        line_lower = line.lower()
        if any(ind in line_lower for ind in company_indicators):
            # Found company context, extract company name
            for offset in range(-1, 2):
                idx = j + offset
                if 0 <= idx < len(lines):
                    candidate = lines[idx].strip()
                    if candidate and len(candidate) < 100 and re.match(r'^[A-Z]', candidate):
                        return candidate
    
    # Strategy 3: Extract first capitalized line in experience section
    for j in range(exp_idx + 1, min(exp_idx + 15, len(lines))):
        line = lines[j].strip()
        if line and len(line) < 100 and re.match(r'^[A-Z]', line) and not line.isupper():
            if len(line.split()) <= 6:
                return line
    
    return None

def extract_designation(text):
    """Extract current designation - search near Present or most recent job"""
    text_lower = text.lower()
    lines = text.split('\n')
    
    # Comprehensive job title keywords
    titles = [
        'engineer', 'developer', 'programmer', 'coder', 'software', 'quality assurance',
        'manager', 'lead', 'leader', 'head', 'director', 'senior', 'principal',
        'junior', 'associate', 'analyst', 'consultant', 'architect', 'specialist',
        'designer', 'scientist', 'research', 'coordinator', 'officer',
        'executive', 'supervisor', 'intern', 'trainee', 'administrator'
    ]
    
    # Find experience section
    exp_idx = -1
    for i, line in enumerate(lines):
        if 'experience' in line.lower() or 'work' in line.lower():
            exp_idx = i
            break
    
    if exp_idx == -1:
        exp_idx = 0
    
    # Strategy 1: Find "Present" and extract designation from nearby lines
    for j in range(exp_idx, min(exp_idx + 30, len(lines))):
        line_lower = lines[j].lower()
        if 'present' in line_lower or 'current' in line_lower:
            # Extract designation from surrounding context
            for offset in range(-2, 3):
                idx = j + offset
                if 0 <= idx < len(lines):
                    candidate = lines[idx].strip()
                    candidate_lower = candidate.lower()
                    if any(title in candidate_lower for title in titles):
                        if len(candidate) < 100:
                            return candidate
    
    # Strategy 2: Search for job title keywords in experience section
    for j in range(exp_idx, min(exp_idx + 20, len(lines))):
        line = lines[j].strip()
        line_lower = line.lower()
        
        if any(title in line_lower for title in titles):
            if 0 < len(line) < 100 and not line.isupper():
                return line
    
    return None

def extract_city(text):
    """Extract city - handle variations, abbreviations, different position indicators"""
    text_lower = text.lower()
    
    # Comprehensive list with aliases and variations
    cities = {
        'mumbai': ['mumbai', 'bombay', 'mum'],
        'delhi': ['delhi', 'new delhi', 'dli'],
        'bangalore': ['bangalore', 'bengaluru', 'blr'],
        'hyderabad': ['hyderabad', 'hyd'],
        'chennai': ['chennai', 'madras', 'chn'],
        'kolkata': ['kolkata', 'calcutta', 'kol'],
        'pune': ['pune', 'poona', 'pun'],
        'ahmedabad': ['ahmedabad', 'ahd'],
        'jaipur': ['jaipur', 'jai'],
        'lucknow': ['lucknow', 'lucknow'],
        'kanpur': ['kanpur', 'kan'],
        'nagpur': ['nagpur', 'nag'],
        'indore': ['indore', 'ind'],
        'bhopal': ['bhopal', 'bho'],
        'visakhapatnam': ['visakhapatnam', 'vizag'],
        'patna': ['patna', 'pat'],
        'vadodara': ['vadodara', 'baroda', 'vad'],
        'ghaziabad': ['ghaziabad', 'gha'],
        'ludhiana': ['ludhiana', 'lud'],
        'agra': ['agra'],
        'nashik': ['nashik'],
        'faridabad': ['faridabad'],
        'meerut': ['meerut'],
        'rajkot': ['rajkot'],
        'guwahati': ['guwahati', 'assam'],
        'chandigarh': ['chandigarh', 'chi'],
        'kochi': ['kochi', 'cochin'],
        'coimbatore': ['coimbatore', 'cbe']
    }
    
    # Search for city mentions
    for city, aliases in cities.items():
        for alias in aliases:
            if alias in text_lower:
                return city.title()
    
    return None

def extract_state(text):
    """Extract state - handle abbreviations and variations"""
    text_lower = text.lower()
    
    # States with aliases and abbreviations
    states = {
        'maharashtra': ['maharashtra', 'mh', 'maharashtra', 'maha'],
        'karnataka': ['karnataka', 'ka', 'karnataka'],
        'tamil nadu': ['tamil nadu', 'tn', 'tamilnadu'],
        'delhi': ['delhi', 'dl', 'new delhi'],
        'uttar pradesh': ['uttar pradesh', 'up', 'up'],
        'west bengal': ['west bengal', 'wb', 'westbengal'],
        'gujarat': ['gujarat', 'gj', 'guj'],
        'rajasthan': ['rajasthan', 'rj', 'raj'],
        'madhya pradesh': ['madhya pradesh', 'mp'],
        'telangana': ['telangana', 'tg', 'ts'],
        'andhra pradesh': ['andhra pradesh', 'ap', 'ap'],
        'kerala': ['kerala', 'kl'],
        'punjab': ['punjab', 'pb', 'pun'],
        'haryana': ['haryana', 'hr', 'har'],
        'bihar': ['bihar', 'br', 'bih'],
        'assam': ['assam', 'as', 'ass'],
        'odisha': ['odisha', 'orissa', 'or'],
        'jharkhand': ['jharkhand', 'jh', 'jhar'],
        'chhattisgarh': ['chhattisgarh', 'ct'],
        'uttarakhand': ['uttarakhand', 'uk', 'uttar']
    }
    
    for state, aliases in states.items():
        for alias in aliases:
            if alias in text_lower:
                return state.title()
    
    return None

if __name__ == '__main__':
    # Production-ready configuration
    debug_mode = os.getenv('FLASK_ENV', 'production') == 'development'
    port = int(os.getenv('PORT', '5050'))
    host = os.getenv('HOST', '127.0.0.1')  # Default to localhost, not 0.0.0.0
    
    # Warn if running in debug mode
    if debug_mode:
        print("\n⚠️  WARNING: Running in DEBUG mode. Do NOT use in production!")
    
    print("\n" + "="*70)
    print("🚀 Resume Parser AI Agent - Production Edition")
    print("="*70)
    print(f"\n📍 Server URL: http://{host}:{port}")
    print(f"\n⚙️  Environment: {'DEVELOPMENT (Debug ON)' if debug_mode else 'PRODUCTION (Debug OFF)'}")
    print(f"🔒 Security: CORS enabled, Headers protected, Input validated")
    print("\n✨ FEATURES:")
    print("   • 🌓 Dark/Light theme toggle")
    print("   • 📊 Real-time parsing progress with individual file tracking")
    print("   • ✅ Success/Failed/Incomplete status indicators")
    print("   • 🔍 Search and filter parsed results")
    print("   • 📥 Export to Excel AND JSON")
    print("   • ⚡ Live statistics dashboard")
    print("   • 🎨 Beautiful modern UI with animations")
    print("   • 🔐 Enterprise-grade security")
    print("\n⚠️  Press CTRL+C to stop the server")
    print("="*70 + "\n")
    
    app.run(debug=debug_mode, host=host, port=port, threaded=True)